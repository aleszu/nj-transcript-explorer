#!/usr/bin/env python3
"""
Reads Word documents, separates by candidate name, and combines them into single documents.
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches

def convert_to_sentence_case(text):
    """
    Convert text to sentence case if it's all caps.
    """
    # Check if text is all caps (ignoring punctuation and spaces)
    words = text.split()
    if not words:
        return text
    
    # Count words that are all caps
    caps_words = sum(1 for word in words if word.isupper() and word.isalpha())
    total_words = sum(1 for word in words if word.isalpha())
    
    # If more than 80% of words are all caps, convert to sentence case
    if total_words > 0 and caps_words / total_words > 0.8:
        # Convert to sentence case
        sentences = text.split('. ')
        converted_sentences = []
        
        for sentence in sentences:
            if sentence.strip():
                # Convert first letter to uppercase, rest to lowercase
                sentence = sentence.strip()
                if sentence:
                    sentence = sentence[0].upper() + sentence[1:].lower()
                converted_sentences.append(sentence)
        
        return '. '.join(converted_sentences)
    
    return text

def read_docx_text(file_path):
    """
    Read all text from a .docx file.
    Equivalent to the read_docx_text function in R.
    """
    try:
        doc = Document(file_path)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                text = paragraph.text.strip()
                # Convert to sentence case if needed
                text = convert_to_sentence_case(text)
                text_parts.append(text)
        
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"Error reading {file_path}: {str(e)}")
        return ""

def extract_date_from_filename(filename):
    """
    Extract date from filename for sorting.
    Looks for patterns like "Oct 2", "Sep 29", "Aug 28", etc.
    """
    import re
    from datetime import datetime
    
    # Month abbreviations mapping
    month_map = {
        'jan': '01', 'feb': '02', 'mar': '03', 'apr': '04',
        'may': '05', 'jun': '06', 'jul': '07', 'aug': '08',
        'sep': '09', 'oct': '10', 'nov': '11', 'dec': '12'
    }
    
    # Look for date patterns like "Oct 2", "Sep 29", "Aug 28"
    date_pattern = r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\s+(\d+)'
    match = re.search(date_pattern, filename.lower())
    
    if match:
        month_abbr = match.group(1)
        day = match.group(2).zfill(2)
        month_num = month_map.get(month_abbr, '12')
        # Assume current year (2025) for sorting
        return f"2025-{month_num}-{day}"
    
    # If no date found, return a very early date for sorting
    return "2025-01-01"

def separate_files_by_candidate(docx_files):
    """
    Separate files by candidate name and sort by date (most recent first).
    Equivalent to the str_detect operations in R.
    """
    sherrill_files = []
    ciattarelli_files = []
    
    for file in docx_files:
        filename_lower = file.name.lower()
        if "sherrill" in filename_lower:
            sherrill_files.append(file)
        elif "ciattarelli" in filename_lower:
            ciattarelli_files.append(file)
    
    # Sort files by date in reverse chronological order (most recent first)
    sherrill_files.sort(key=lambda x: extract_date_from_filename(x.name), reverse=True)
    ciattarelli_files.sort(key=lambda x: extract_date_from_filename(x.name), reverse=True)
    
    return sherrill_files, ciattarelli_files

def combine_text_from_files(file_list, separator="\n\n-----\n\n"):
    """
    Read and combine text from a list of files.
    Equivalent to map_chr and paste operations in R.
    """
    combined_texts = []
    
    for file_path in file_list:
        text = read_docx_text(file_path)
        if text:  # Only add non-empty text
            combined_texts.append(text)
    
    return separator.join(combined_texts)

def create_combined_document(text_content, base_filename):
    """
    Create a new Word document with combined text.
    Equivalent to read_docx() %>% body_add_par() %>% print() in R.
    """
    from datetime import datetime
    
    try:
        doc = Document()
        
        # Add the combined text as a single paragraph
        # Split by separator to create separate paragraphs for each file
        paragraphs = text_content.split("\n\n-----\n\n")
        
        for i, paragraph_text in enumerate(paragraphs):
            if paragraph_text.strip():
                # Add a heading for each file section
                if i > 0:
                    doc.add_heading(f"File {i+1}", level=2)
                
                # Add the paragraph content
                para = doc.add_paragraph(paragraph_text)
        
        # Add today's date to filename
        today = datetime.now()
        date_str = today.strftime("%m%d%Y")  # MMDDYYYY format
        output_filename = f"{base_filename}_{date_str}.docx"
        
        # Save the document
        doc.save(output_filename)
        print(f"Created: {output_filename}")
        
    except Exception as e:
        print(f"Error creating {output_filename}: {str(e)}")

def main():
    """
    Main function to process Word documents and create combined files.
    """
    # Set working directory (equivalent to setwd() in R)
    working_dir = Path.cwd()  # Current directory
    print(f"Working directory: {working_dir}")
    
    # List all .docx files in the folder (equivalent to list.files() in R)
    docx_files = list(working_dir.glob("*.docx"))
    # Exclude temporary files, requirements, output files, and NJ Governors Forum docs
    docx_files = [f for f in docx_files if not f.name.startswith("~$") 
                  and f.name != "requirements.docx"
                  and not f.name.startswith("Combined_")
                  and "NJGovernorsForum" not in f.name]
    
    if not docx_files:
        print("No .docx files found in the current directory!")
        return
    
    print(f"Found {len(docx_files)} Word documents:")
    for file in docx_files:
        print(f"  - {file.name}")
    
    # Separate files by candidate name
    sherrill_files, ciattarelli_files = separate_files_by_candidate(docx_files)
    
    print(f"\nSherrill files ({len(sherrill_files)}) - sorted by date (most recent first):")
    for file in sherrill_files:
        print(f"  - {file.name}")
    
    print(f"\nCiattarelli files ({len(ciattarelli_files)}) - sorted by date (most recent first):")
    for file in ciattarelli_files:
        print(f"  - {file.name}")
    
    # Process Sherrill files
    if sherrill_files:
        print("\nProcessing Sherrill files...")
        sherrill_text = combine_text_from_files(sherrill_files)
        create_combined_document(sherrill_text, "Combined_Sherrill_Text")
    
    # Process Ciattarelli files
    if ciattarelli_files:
        print("\nProcessing Ciattarelli files...")
        ciattarelli_text = combine_text_from_files(ciattarelli_files)
        create_combined_document(ciattarelli_text, "Combined_Ciattarelli_Text")
    
    print("\nProcessing complete!")

if __name__ == "__main__":
    main()
