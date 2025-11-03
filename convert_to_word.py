#!/usr/bin/env python3
"""
Script to convert transcript text files to Word documents.
Looks for files in the /transcripts directory.
Outputs to the /data directory.
Processes files with naming convention: Candidate_Location_Date.txt
Creates Word documents with header: Date - Candidate - Location
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_filename(filename):
    """
    Parse filename to extract candidate, location, and date.
    Expected format: Candidate_Location_Date.txt
    """
    # Remove .txt extension
    name_without_ext = filename.replace('.txt', '')
    
    # Split by underscore
    parts = name_without_ext.split('_')
    
    if len(parts) >= 3:
        candidate = parts[0]
        location = parts[1]
        date = parts[2]
        
        # Format date (assuming format like 10022025)
        if len(date) == 8:
            month_num = int(date[:2])
            day = int(date[2:4])
            year = date[4:]
            
            # Convert month number to abbreviation
            month_abbr = {
                1: 'Jan', 2: 'Feb', 3: 'Mar', 4: 'Apr',
                5: 'May', 6: 'Jun', 7: 'Jul', 8: 'Aug',
                9: 'Sep', 10: 'Oct', 11: 'Nov', 12: 'Dec'
            }.get(month_num, 'Unknown')
            
            formatted_date = f"{month_abbr} {day}"
        else:
            formatted_date = date
            
        return candidate, location, formatted_date
    else:
        # Fallback if format doesn't match expected pattern
        return "Unknown", "Unknown", "Unknown"

def remove_line_breaks(text):
    """
    Remove line breaks and clean up text formatting.
    """
    # Replace line breaks with spaces
    text = text.replace('\n', ' ')
    
    # Clean up multiple spaces
    text = re.sub(r'\s+', ' ', text)
    
    # Clean up HTML entities
    text = text.replace('&#39;', "'")
    text = text.replace('&quot;', '"')
    
    return text.strip()

def create_word_document(candidate, location, date, content, base_filename, output_dir):
    """
    Create a Word document with the specified header and content.
    
    Args:
        candidate: Candidate name
        location: Location/event name
        date: Date string
        content: Document content
        base_filename: Base filename for output
        output_dir: Directory to save the document
    """
    doc = Document()
    
    # Add header: Date - Candidate - Location
    header = doc.add_heading(f"{date} - {candidate} - {location}", level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some spacing
    doc.add_paragraph()
    
    # Add the content
    content_para = doc.add_paragraph(content)
    
    # Ensure output directory exists
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Create output filename with date prefix in the output directory
    output_filename = output_dir / f"{date} - {base_filename}"
    
    # Save the document (will overwrite if exists)
    doc.save(str(output_filename))
    print(f"Created: {output_filename}")

def process_text_files(directory_path, output_dir):
    """
    Process all .txt files in the specified directory.
    
    Args:
        directory_path: Directory containing .txt files to process
        output_dir: Directory to save output DOCX files
    """
    directory = Path(directory_path)
    
    if not directory.exists():
        print(f"Directory {directory_path} does not exist!")
        return
    
    # Find all .txt files (excluding requirements.txt)
    txt_files = [f for f in directory.glob("*.txt") if f.name != "requirements.txt"]
    
    if not txt_files:
        print(f"No .txt files found in {directory_path}")
        return
    
    print(f"Found {len(txt_files)} text files to process:")
    print(f"Output directory: {output_dir}")
    
    for txt_file in txt_files:
        print(f"\nProcessing: {txt_file.name}")
        
        try:
            # Parse filename
            candidate, location, date = parse_filename(txt_file.name)
            print(f"  Candidate: {candidate}")
            print(f"  Location: {location}")
            print(f"  Date: {date}")
            
            # Read file content
            with open(txt_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Remove line breaks
            cleaned_content = remove_line_breaks(content)
            
            # Create base filename (without date prefix)
            base_filename = txt_file.with_suffix('.docx').name
            
            # Create Word document (will overwrite if exists)
            create_word_document(candidate, location, date, cleaned_content, base_filename, output_dir)
            
        except Exception as e:
            print(f"  Error processing {txt_file.name}: {str(e)}")

def main():
    """
    Main function to run the script.
    """
    # Get the directory where the script is located
    script_dir = Path(__file__).parent
    
    # Set input directory to /transcripts (relative to script location)
    directory_path = script_dir / 'transcripts'
    
    # Set output directory to /app/data (relative to script location)
    output_dir = script_dir / 'data'
    
    print("Transcript to Word Converter")
    print("=" * 40)
    print(f"Processing files in: {directory_path}")
    print(f"Output directory: {output_dir}")
    print("Note: Existing files will be overwritten")
    print("=" * 40)
    
    process_text_files(directory_path, output_dir)
    
    print("\n" + "=" * 40)
    print("Processing complete!")

if __name__ == "__main__":
    main()
